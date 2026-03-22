"""
Document parsing service for multiple file formats
多文件格式的文档解析服务

Day 2 Enhancement: Supports PDF, Word, HTML, Markdown, TXT
Day 2 增强： 支持 PDF, Word, HTML, Markdown, TXT
"""

import os
import tempfile
from typing import List, Dict, Optional, Tuple
from datetime import datetime

# PDF parsing
# PDF 解析
from pypdf import PdfReader

# Word parsing
# Word 解析
from docx import Document

# HTML parsing
# HTML 解析
from bs4 import BeautifulSoup
import lxml.html

# Encoding detection
# 编码检测
import chardet

# LangChain text splitters
# LangChain 文本分割器
from langchain_text_splitters import (
    RecursiveCharacterTextSplitter,
    MarkdownHeaderTextSplitter,
    HTMLHeaderTextSplitter,
)

from config import settings


class DocumentInfo:
    """
    Document metadata information
    文档元数据信息
    """
    def __init__(
        self,
        filename: str,
        file_type: str,
        file_size: int,
        created_at: datetime,
        metadata: Optional[Dict] = None
    ):
        self.filename = filename
        self.file_type = file_type
        self.file_size = file_size
        self.created_at = created_at
        self.metadata = metadata or {}


class ParsedDocument:
    """
    Parsed document with content and chunks
    解析后的文档及其内容和分块
    """
    def __init__(
        self,
        content: str,
        chunks: List[str],
        metadata: Dict,
        document_info: DocumentInfo
    ):
        self.content = content
        self.chunks = chunks
        self.metadata = metadata
        self.document_info = document_info


class DocumentParserService:
    """
    Service for parsing various document formats
    解析各种文档格式的服务
    """

    # Supported file extensions
    # 支持的文件扩展名
    SUPPORTED_EXTENSIONS = {
        '.txt': 'text',
        '.md': 'markdown',
        '.pdf': 'pdf',
        '.docx': 'word',
        '.doc': 'word',
        '.html': 'html',
        '.htm': 'html',
    }

    def __init__(self):
        """
        Initialize the document parser service
        初始化文档解析服务
        """
        # Initialize text splitters
        # 初始化文本分割器
        self._text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            separators=["\n\n", "\n", "。", ".", " ", ""],
        )

        self._markdown_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=["#", "##", "###"],
        )

        self._html_splitter = HTMLHeaderTextSplitter(
            headers_to_split_on=[("h1", "Header 1"), ("h2", "Header 2"), ("h3", "Header 3")]
        )

    def get_file_type(self, filename: str) -> Optional[str]:
        """
        Get the file type from filename extension
        根据文件扩展名获取文件类型

        Args:
            filename: Name of the file
                     文件名
        Returns:
            File type or None if unsupported
            文件类型，如果不支持则返回 None
        """
        ext = os.path.splitext(filename)[1].lower()
        return self.SUPPORTED_EXTENSIONS.get(f'.{ext}')

    def is_supported(self, filename: str) -> bool:
        """
        Check if file type is supported
        检查文件类型是否支持

        Args:
            filename: Name of the file
                     文件名
        Returns:
            True if supported, False otherwise
            如果支持返回 True，否则返回 False
        """
        return self.get_file_type(filename) is not None

    def detect_encoding(self, file_content: bytes) -> str:
        """
        Detect the encoding of file content
        检测文件内容的编码

        Args:
            file_content: Raw file bytes
                          原始文件字节
        Returns:
            Detected encoding
            检测到的编码
        """
        result = chardet.detect(file_content)
        return result['encoding'] or 'utf-8'

    async def parse_file(
        self,
        file_content: bytes,
        filename: str
    ) -> ParsedDocument:
        """
        Parse a file and return content, chunks, and metadata
        解析文件并返回内容、分块和元数据

        Args:
            file_content: Raw file bytes
                          原始文件字节
            filename: Original filename
                       原始文件名
        Returns:
            ParsedDocument with content, chunks, metadata
            包含内容、分块、元数据的 ParsedDocument
        """
        file_type = self.get_file_type(filename)
        if not file_type:
            raise ValueError(f"Unsupported file type: {filename}")

        # Extract metadata
        # 提取元数据
        metadata = await self._extract_metadata(file_content, filename, file_type)

        # Parse content based on file type
        # 根据文件类型解析内容
        if file_type == 'text':
            content = await self._parse_text(file_content, filename)
        elif file_type == 'markdown':
            content = await self._parse_markdown(file_content, filename)
        elif file_type == 'pdf':
            content = await self._parse_pdf(file_content, filename)
        elif file_type == 'word':
            content = await self._parse_word(file_content, filename)
        elif file_type == 'html':
            content = await self._parse_html(file_content, filename)
        else:
            raise ValueError(f"Unknown file type: {file_type}")

        # Split content into chunks
        # 将内容分割为分块
        chunks = await self._split_content(content, file_type)

        # Create document info
        # 创建文档信息
        doc_info = DocumentInfo(
            filename=filename,
            file_type=file_type,
            file_size=len(file_content),
            created_at=datetime.now(),
            metadata=metadata
        )

        return ParsedDocument(
            content=content,
            chunks=chunks,
            metadata=metadata,
            document_info=doc_info
        )

    async def _extract_metadata(
        self,
        file_content: bytes,
        filename: str,
        file_type: str
    ) -> Dict:
        """
        Extract metadata from file
        从文件中提取元数据

        Args:
            file_content: Raw file bytes
                          原始文件字节
            filename: Original filename
                       原始文件名
            file_type: Type of the file
                        文件类型
        Returns:
            Dictionary of metadata
            元数据字典
        """
        metadata = {
            'filename': filename,
            'file_type': file_type,
            'file_size': len(file_content),
        }

        # Try to extract title from content
        # 尝试从内容中提取标题
        try:
            encoding = self.detect_encoding(file_content)
            text_preview = file_content[:2000].decode(encoding, errors='ignore')

            # Look for title patterns
            # 查找标题模式
            lines = text_preview.split('\n')[:10]
            for line in lines:
                line = line.strip()
                if line and len(line) < 100:
                    # Skip markdown heading markers
                    # 跳过 markdown 标题标记
                    clean_line = line.lstrip('#').strip()
                    if clean_line and len(clean_line) > 3:
                        metadata['title'] = clean_line
                        break
        except Exception:
            pass

        return metadata

    async def _parse_text(self, file_content: bytes, filename: str) -> str:
        """
        Parse plain text file
        解析纯文本文件

        Args:
            file_content: Raw file bytes
                          原始文件字节
            filename: Original filename
                       原始文件名
        Returns:
            Text content
            文本内容
        """
        encoding = self.detect_encoding(file_content)
        return file_content.decode(encoding)

    async def _parse_markdown(self, file_content: bytes, filename: str) -> str:
        """
        Parse Markdown file
        解析 Markdown 文件

        Args:
            file_content: Raw file bytes
                          原始文件字节
            filename: Original filename
                       原始文件名
        Returns:
            Text content (Markdown preserved)
            文本内容（保留 Markdown 格式）
        """
        encoding = self.detect_encoding(file_content)
        return file_content.decode(encoding)

    async def _parse_pdf(self, file_content: bytes, filename: str) -> str:
        """
        Parse PDF file using PyPDF2
        使用 PyPDF2 解析 PDF 文件

        Args:
            file_content: Raw file bytes
                          原始文件字节
            filename: Original filename
                       原始文件名
        Returns:
            Extracted text content
            提取的文本内容
        """
        text_parts = []

        # Write to temporary file for PyPDF2
        # 写入临时文件供 PyPDF2 使用
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_file:
            tmp_file.write(file_content)
            tmp_file.flush()

            try:
                reader = PdfReader(tmp_file.name)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            finally:
                # Clean up temp file
                # 清理临时文件
                try:
                    os.unlink(tmp_file.name)
                except:
                    pass

        return '\n\n'.join(text_parts)

    async def _parse_word(self, file_content: bytes, filename: str) -> str:
        """
        Parse Word file using python-docx
        使用 python-docx 解析 Word 文件

        Args:
            file_content: Raw file bytes
                          原始文件字节
            filename: Original filename
                       原始文件名
        Returns:
            Extracted text content
            提取的文本内容
        """
        text_parts = []

        # Write to temporary file for python-docx
        # 写入临时文件供 python-docx 使用
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            tmp_file.write(file_content)
            tmp_file.flush()

            try:
                doc = Document(tmp_file.name)
                for paragraph in doc.paragraphs:
                    para_text = paragraph.text
                    if para_text:
                        text_parts.append(para_text)

                # Also extract text from tables
                # 同时从表格中提取文本
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text:
                                text_parts.append(cell.text)
            finally:
                # Clean up temp file
                # 清理临时文件
                try:
                    os.unlink(tmp_file.name)
                except:
                    pass

        return '\n\n'.join(text_parts)

    async def _parse_html(self, file_content: bytes, filename: str) -> str:
        """
        Parse HTML file using BeautifulSoup
        使用 BeautifulSoup 解析 HTML 文件

        Args:
            file_content: Raw file bytes
                          原始文件字节
            filename: Original filename
                        原始文件名
        Returns:
            Extracted text content
            提取的文本内容
        """
        encoding = self.detect_encoding(file_content)
        soup = BeautifulSoup(file_content, 'lxml')

        # Remove script and style elements
        # 移除脚本和样式元素
        for script in soup(['script', 'style']):
            script.decompose()

        # Get text
        # 获取文本
        text = soup.get_text(separator='\n')

        # Clean up whitespace
        # 清理空白
        lines = [line.strip() for line in text.split('\n')]
        lines = [line for line in lines if line]

        return '\n\n'.join(lines)

    async def _split_content(self, content: str, file_type: str) -> List[str]:
        """
        Split content into chunks using appropriate strategy
        使用适当的策略将内容分割为分块

        Args:
            content: Text content to split
                     要分割的文本内容
            file_type: Type of the file
                       文件类型
        Returns:
            List of text chunks
            文本分块列表
        """
        if not content:
            return []

        # Use appropriate splitter based on file type
        # 根据文件类型使用适当的分割器
        if file_type == 'markdown':
            try:
                # Try markdown-specific splitting first
                # 首先尝试 Markdown 特定的分割
                md_chunks = self._markdown_splitter.split_text(content)
                # If chunks are too large, use recursive splitter
                # 如果分块太大，使用递归分割器
                final_chunks = []
                for chunk in md_chunks:
                    if len(chunk.page_content) > settings.chunk_size * 2:
                        sub_chunks = self._text_splitter.split_text(chunk.page_content)
                        final_chunks.extend(sub_chunks)
                    else:
                        final_chunks.append(chunk.page_content)
                return final_chunks
            except Exception:
                # Fall back to recursive splitter
                # 回退到递归分割器
                return self._text_splitter.split_text(content)

        elif file_type == 'html':
            try:
                # Try HTML-specific splitting first
                # 首先尝试 HTML 特定的分割
                html_chunks = self._html_splitter.split_text(content)
                # If chunks are too large, use recursive splitter
                # 如果分块太大，使用递归分割器
                final_chunks = []
                for chunk in html_chunks:
                    if len(chunk.page_content) > settings.chunk_size * 2:
                        sub_chunks = self._text_splitter.split_text(chunk.page_content)
                        final_chunks.extend(sub_chunks)
                    else:
                        final_chunks.append(chunk.page_content)
                return final_chunks
            except Exception:
                # Fall back to recursive splitter
                # 回退到递归分割器
                return self._text_splitter.split_text(content)

        else:
            # Use recursive character splitter for text, PDF, Word
            # 对文本、PDF、Word 使用递归字符分割器
            return self._text_splitter.split_text(content)


# Global document parser service instance
# 全局文档解析服务实例
document_parser = DocumentParserService()
